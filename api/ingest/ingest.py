#!/usr/bin/env python3
"""
⚠️ TEMPORARY INGESTION SCRIPT ⚠️

Purpose:
    - Scrape and ingest university regulation documents from Bilgi University
    - Downloads PDFs, DOCX, and HTML pages
    - Extracts text content and generates embeddings
    - Populates the university_documents table

Usage:
    cd api/database/ingest
    python3 ingest_regulations.py

Requirements:
    - PostgreSQL with pgvector extension running
    - Embedding service running (docker-compose up -d embeddings)
    - pip install requests pypdf python-docx beautifulsoup4
"""

import json
import sys
import os
import io
import requests
from pathlib import Path
from sqlalchemy import create_engine, Column, Integer, String, Text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from pgvector.sqlalchemy import Vector

# Add parent directory to path to import services
sys.path.insert(0, str(Path(__file__).parent.parent.parent))
from services.embedding_service import get_embedding_service

# Database connection
DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://user:password@localhost:5432/orbisdb")

engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(bind=engine)

Base = declarative_base()

# Initialize embedding service
print("Initializing embedding service...")
try:
    embedding_service = get_embedding_service()
    EMBEDDING_DIM = embedding_service.get_dimension()
    print(f"✓ Embedding service ready! Dimension: {EMBEDDING_DIM}")
except Exception as e:
    print(f"❌ Failed to initialize embedding service: {e}")
    print("Make sure Docker container is running: docker-compose up -d embeddings")
    sys.exit(1)


class UniversityDocument(Base):
    """
    Represents a raw, unchunked document scraped from a university resource.
    """
    __tablename__ = "university_documents"
    
    id = Column(Integer, primary_key=True)
    source_url = Column(String(500), unique=True, nullable=False)
    title = Column(String(255), nullable=False)
    raw_content = Column(Text, nullable=False)
    summary = Column(Text)
    keywords = Column(Text)
    keyword_embedding = Column(Vector(EMBEDDING_DIM))


def fetch_content(url: str) -> str:
    """
    Fetch and extract text content from a URL.
    Handles PDF, DOCX, and HTML formats.
    """
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    }
    
    print(f"  Fetching: {url}")
    response = requests.get(url, headers=headers, timeout=30)
    response.raise_for_status()
    
    content_type = response.headers.get('Content-Type', '').lower()
    
    # Handle PDF
    if 'pdf' in content_type or url.endswith('.pdf'):
        return extract_pdf_text(response.content)
    
    # Handle DOCX
    if 'wordprocessingml' in content_type or url.endswith('.docx'):
        return extract_docx_text(response.content)
    
    # Handle HTML (default)
    return extract_html_text(response.text)


def extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF binary content."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text() or '')
        return '\n'.join(text_parts)
    except ImportError:
        print("    ⚠️ pypdf not installed. Install with: pip install pypdf")
        raise
    except Exception as e:
        print(f"    ⚠️ PDF extraction error: {e}")
        raise


def extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX binary content."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        return '\n'.join(text_parts)
    except ImportError:
        print("    ⚠️ python-docx not installed. Install with: pip install python-docx")
        raise
    except Exception as e:
        print(f"    ⚠️ DOCX extraction error: {e}")
        raise


def extract_html_text(html: str) -> str:
    """Extract text from HTML content."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        
        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'header', 'footer']):
            element.decompose()
        
        # Get text
        text = soup.get_text(separator='\n')
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        return '\n'.join(chunk for chunk in chunks if chunk)
    except ImportError:
        print("    ⚠️ beautifulsoup4 not installed. Install with: pip install beautifulsoup4")
        raise


def ingest_regulations(json_filepath: str):
    """Load regulation documents from JSON metadata file and ingest into database."""
    
    # Load metadata
    with open(json_filepath, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    documents = data.get('documents', data) if isinstance(data, dict) else data
    
    db = SessionLocal()
    
    try:
        for doc_meta in documents:
            source_url = doc_meta['source_url']
            title = doc_meta['title']
            summary = doc_meta.get('summary', '')
            keywords = doc_meta.get('keywords', '')
            
            print(f"\nProcessing: {title}")
            
            # Check if already exists
            existing = db.query(UniversityDocument).filter_by(source_url=source_url).first()
            if existing:
                print(f"  ⚠️ Already exists, skipping...")
                continue
            
            # Fetch and extract content
            try:
                raw_content = fetch_content(source_url)
                if not raw_content or len(raw_content.strip()) < 100:
                    print(f"  ⚠️ Content too short or empty, skipping...")
                    continue
            except Exception as e:
                print(f"  ❌ Failed to fetch content: {e}")
                continue
            
            # Generate embedding from keywords
            embed_text = f"{title}. {summary}. Keywords: {keywords}"
            print(f"  Generating embedding...")
            keyword_embedding = embedding_service.embed_text(embed_text)
            
            # Create document record
            document = UniversityDocument(
                source_url=source_url,
                title=title,
                raw_content=raw_content,
                summary=summary,
                keywords=keywords,
                keyword_embedding=keyword_embedding
            )
            
            db.add(document)
            db.flush()
            
            print(f"  ✓ Added: {title} ({len(raw_content)} chars)")
        
        db.commit()
        print(f"\n✓ Successfully imported regulations with embeddings!")
        
    except Exception as e:
        db.rollback()
        print(f"\n✗ Error: {e}")
        raise
    finally:
        db.close()


# Regulation document metadata
REGULATIONS_METADATA = [
    {
        "source_url": "https://www.bilgi.edu.tr/upload/credit-system-bachelors-degree-and-associate-degree-education-and-examination-regulation/",
        "title": "Credit System Undergraduate and Two-Year Program Education and Examination Regulation",
        "summary": "Comprehensive regulation governing undergraduate and associate degree education at Istanbul Bilgi University, covering admission, registration, course credits, examinations, grading, graduation requirements, and academic standing.",
        "keywords": "undergraduate, associate degree, credit system, examination, registration, tuition, grading, GPA, graduation, academic standing, courses, diploma"
    },
    {
        "source_url": "https://www.bilgi.edu.tr/media/uploads/2024/03/26/regulation-on-double-major-minor-and-honors-programs.pdf",
        "title": "Regulation on Double Major, Minor and Honors Programs",
        "summary": "Regulation defining requirements and procedures for double major, minor, and honors programs including application conditions, credit loads, academic standing requirements, and graduation criteria.",
        "keywords": "double major, minor program, honors program, secondary major, GPA requirements, graduation, credit load, capstone thesis"
    },
    {
        "source_url": "https://www.bilgi.edu.tr/media/uploads/2018/08/08/academic-advisor-directive.pdf",
        "title": "Academic Advisor Directive",
        "summary": "Directive defining academic advising responsibilities for instructors and student obligations during the advising process, including advisor assignment procedures and course registration guidance.",
        "keywords": "academic advisor, course selection, student guidance, registration, advisor responsibilities, student responsibilities"
    },
    {
        "source_url": "https://www.bilgi.edu.tr/media/uploads/2022/03/09/financialproceduresandprinciplesforstudentsundergraduateandassociatedegree.pdf",
        "title": "Financial Procedures and Principles for Students (Undergraduate and Associate Degree)",
        "summary": "Financial regulations covering tuition fees, enrollment suspension and cancellation procedures, refund policies, and payment obligations for undergraduate and associate degree students.",
        "keywords": "tuition fee, enrollment suspension, enrollment cancellation, refund, payment, financial procedures, scholarship, late enrollment"
    },
    {
        "source_url": "https://www.bilgi.edu.tr/media/uploads/2024/11/26/graduate-education-and-training-regulations.docx",
        "title": "Graduate Education and Training Regulations",
        "summary": "Comprehensive regulation governing master's and doctoral programs including admission, course requirements, thesis/project procedures, examinations, and graduation requirements.",
        "keywords": "graduate education, masters degree, doctoral program, thesis, dissertation, graduate courses, PhD, academic requirements"
    },
    {
        "source_url": "https://www.bilgi.edu.tr/media/uploads/2022/03/09/financialproceduresandprinciplesforstudentsgraduatedegree.pdf",
        "title": "Financial Procedures and Principles for Students (Graduate Degree)",
        "summary": "Financial regulations covering tuition fees, enrollment suspension and cancellation procedures, refund policies, and payment obligations specifically for graduate degree students.",
        "keywords": "graduate tuition, enrollment suspension, enrollment cancellation, refund, payment, financial procedures, graduate scholarship"
    }
]


def ingest_from_embedded_metadata():
    """Ingest regulations using embedded metadata (no external JSON file needed)."""
    
    db = SessionLocal()
    success_count = 0
    
    try:
        for doc_meta in REGULATIONS_METADATA:
            source_url = doc_meta['source_url']
            title = doc_meta['title']
            summary = doc_meta.get('summary', '')
            keywords = doc_meta.get('keywords', '')
            
            print(f"\nProcessing: {title}")
            
            # Check if already exists
            existing = db.query(UniversityDocument).filter_by(source_url=source_url).first()
            if existing:
                print(f"  ⚠️ Already exists, skipping...")
                continue
            
            # Fetch and extract content
            try:
                raw_content = fetch_content(source_url)
                if not raw_content or len(raw_content.strip()) < 100:
                    print(f"  ⚠️ Content too short or empty, skipping...")
                    continue
            except Exception as e:
                print(f"  ❌ Failed to fetch content: {e}")
                continue
            
            # Generate embedding from keywords
            embed_text = f"{title}. {summary}. Keywords: {keywords}"
            print(f"  Generating embedding...")
            keyword_embedding = embedding_service.embed_text(embed_text)
            
            # Create document record
            document = UniversityDocument(
                source_url=source_url,
                title=title,
                raw_content=raw_content,
                summary=summary,
                keywords=keywords,
                keyword_embedding=keyword_embedding
            )
            
            db.add(document)
            db.flush()
            success_count += 1
            
            print(f"  ✓ Added: {title} ({len(raw_content)} chars)")
        
        db.commit()
        print(f"\n✓ Successfully imported {success_count} regulation documents!")
        
    except Exception as e:
        db.rollback()
        print(f"\n✗ Error: {e}")
        raise
    finally:
        db.close()


if __name__ == "__main__":
    # Create tables if they don't exist
    print("Checking/Creating database tables...")
    Base.metadata.create_all(engine)
    print("✓ Tables ready")

    if len(sys.argv) > 1:

        # Use external JSON file if provided
        ingest_regulations(sys.argv[1])
    else:
        # Use embedded metadata
        ingest_from_embedded_metadata()